"""Report export page for PDF / Excel / Word outputs."""
from __future__ import annotations

import io
from decimal import Decimal, ROUND_HALF_UP
from typing import Dict

import pandas as pd
import streamlit as st
from docx import Document
from fpdf import FPDF

from calc import compute, plan_from_models, summarize_plan_metrics
from formatting import UNIT_FACTORS, format_amount_with_unit, format_ratio, to_decimal
from state import ensure_session_defaults, load_finance_bundle
from theme import inject_theme


PDF_UNIT_LABELS = {
    "百万円": "million",
    "千円": "thousand",
    "円": "",
}


def _ascii_unit_label(unit: str) -> str:
    label = PDF_UNIT_LABELS.get(unit)
    if label is not None:
        return label
    safe = "".join(ch for ch in unit if 32 <= ord(ch) <= 126)
    return safe


def format_amount_for_pdf(value: object, unit: str) -> str:
    try:
        amount = to_decimal(value)
    except Exception:
        return "N/A"
    factor = UNIT_FACTORS.get(unit, Decimal("1"))
    if factor == 0:
        factor = Decimal("1")
    scaled = amount / factor
    if scaled.is_nan() or scaled.is_infinite():
        return "N/A"
    quant = Decimal("1") if abs(scaled) >= 1 else Decimal("0.01")
    scaled = scaled.quantize(quant, rounding=ROUND_HALF_UP)
    formatted = f"{scaled:,.0f}" if quant == Decimal("1") else f"{scaled:,.2f}"
    label = _ascii_unit_label(unit)
    result = f"JPY {formatted}"
    if label:
        result = f"{result} {label}"
    return result.strip()


def format_ratio_for_pdf(value: object) -> str:
    try:
        ratio = to_decimal(value)
    except Exception:
        return "N/A"
    if ratio.is_nan() or ratio.is_infinite():
        return "N/A"
    return f"{ratio * Decimal('100'):.1f}%"


def build_pdf_summary_lines(
    *, fiscal_year: int, unit: str, amounts: Dict[str, Decimal], metrics: Dict[str, Decimal]
) -> list[str]:
    return [
        f"FY{fiscal_year} Plan Summary",
        f"Revenue: {format_amount_for_pdf(amounts.get('REV', Decimal('0')), unit)}",
        f"Gross margin: {format_ratio_for_pdf(metrics.get('gross_margin'))}",
        f"Ordinary profit: {format_amount_for_pdf(amounts.get('ORD', Decimal('0')), unit)}",
        f"Ordinary profit margin: {format_ratio_for_pdf(metrics.get('ord_margin'))}",
        f"Break-even revenue: {format_amount_for_pdf(metrics.get('breakeven', Decimal('0')), unit)}",
    ]

st.set_page_config(
    page_title="経営計画スタジオ｜Report",
    page_icon="📝",
    layout="wide",
)

inject_theme()
ensure_session_defaults()

settings_state: Dict[str, object] = st.session_state.get("finance_settings", {})
unit = str(settings_state.get("unit", "百万円"))
fte = Decimal(str(settings_state.get("fte", 20)))
fiscal_year = int(settings_state.get("fiscal_year", 2025))

bundle, has_custom_inputs = load_finance_bundle()
if not has_custom_inputs:
    st.info("入力データが未保存のため、既定値でレポートを生成します。")

plan_cfg = plan_from_models(
    bundle.sales,
    bundle.costs,
    bundle.capex,
    bundle.loans,
    bundle.tax,
    fte=fte,
    unit=unit,
)

amounts = compute(plan_cfg)
metrics = summarize_plan_metrics(amounts)

st.title("📝 レポート出力")
st.caption("主要指標とKPIのサマリーをPDF / Excel / Word形式でダウンロードできます。")

pdf_tab, excel_tab, word_tab = st.tabs(["PDF", "Excel", "Word"])

summary_lines_ja = [
    f"FY{fiscal_year} 計画サマリー",
    f"売上高: {format_amount_with_unit(amounts.get('REV', Decimal('0')), unit)}",
    f"粗利率: {format_ratio(metrics.get('gross_margin'))}",
    f"経常利益: {format_amount_with_unit(amounts.get('ORD', Decimal('0')), unit)}",
    f"経常利益率: {format_ratio(metrics.get('ord_margin'))}",
    f"損益分岐点売上高: {format_amount_with_unit(metrics.get('breakeven', Decimal('0')), unit)}",
]

pdf_summary_lines = build_pdf_summary_lines(
    fiscal_year=fiscal_year, unit=unit, amounts=amounts, metrics=metrics
)

with pdf_tab:
    st.subheader("PDFレポート")
    pdf_buffer = io.BytesIO()
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=14)
    pdf.cell(0, 10, txt="Keiei Keikaku Studio | Summary Report", ln=True)
    pdf.set_font("Helvetica", size=11)
    for line in pdf_summary_lines:
        pdf.multi_cell(0, 8, line)
    pdf.output(pdf_buffer)
    st.download_button(
        "📄 PDFダウンロード",
        data=pdf_buffer.getvalue(),
        file_name=f"plan_report_{fiscal_year}.pdf",
        mime="application/pdf",
    )

with excel_tab:
    st.subheader("Excelレポート")
    excel_buffer = io.BytesIO()
    with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
        pd.DataFrame(
            {
                "項目": ["売上高", "粗利", "営業利益", "経常利益"],
                "金額": [
                    float(amounts.get("REV", Decimal("0"))),
                    float(amounts.get("GROSS", Decimal("0"))),
                    float(amounts.get("OP", Decimal("0"))),
                    float(amounts.get("ORD", Decimal("0"))),
                ],
            }
        ).to_excel(writer, sheet_name="PL", index=False)

        pd.DataFrame(
            {
                "指標": ["粗利率", "営業利益率", "経常利益率", "損益分岐点"],
                "値": [
                    float(metrics.get("gross_margin", Decimal("0"))),
                    float(metrics.get("op_margin", Decimal("0"))),
                    float(metrics.get("ord_margin", Decimal("0"))),
                    float(metrics.get("breakeven", Decimal("0"))),
                ],
            }
        ).to_excel(writer, sheet_name="KPI", index=False)

    st.download_button(
        "📊 Excelダウンロード",
        data=excel_buffer.getvalue(),
        file_name=f"plan_report_{fiscal_year}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

with word_tab:
    st.subheader("Wordレポート")
    doc = Document()
    doc.add_heading("経営計画スタジオ レポート", level=1)
    doc.add_paragraph(f"FY{fiscal_year} / 表示単位: {unit} / FTE: {fte}")
    doc.add_paragraph("主要KPI")
    bullet = doc.add_paragraph()
    bullet.style = "List Bullet"
    bullet.add_run(summary_lines_ja[1])
    for line in summary_lines_ja[2:]:
        para = doc.add_paragraph()
        para.style = "List Bullet"
        para.add_run(line)

    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    st.download_button(
        "📝 Wordダウンロード",
        data=word_buffer.getvalue(),
        file_name=f"plan_report_{fiscal_year}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
